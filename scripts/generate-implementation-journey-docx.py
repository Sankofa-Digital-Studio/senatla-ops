from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc" / "senatla-ops-implementation-journey.docx"
BLUE = "17365D"
ACCENT = "D6A800"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "59636E"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend([begin, instr, separate, text, end])


def add_page_number(paragraph):
    paragraph.add_run("Page ")
    add_field(paragraph, "PAGE")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:color"), "D7DEE8")
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=BLUE)
    p.add_run(text)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, BLUE)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=9.5, color=WHITE, bold=True)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        row_properties = cells[0]._tc.getparent().get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(row_data):
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, size=9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SENATLA OPS  |  IMPLEMENTATION JOURNEY")
    set_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)
    for run in footer.runs:
        set_font(run, size=8.5, color=MUTED)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WORKING IMPLEMENTATION RECORD")
    set_font(r, size=10, color=ACCENT, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Senatla Ops\nImplementation Journey")
    set_font(r, size=29, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Enhancing the existing operational system through evidence-led vertical slices")
    set_font(r, size=14, color=MUTED, italic=True)
    add_table(doc, ["Document status", "Baseline", "Current branch"], [["Living working document", "24 August 2026", "sankofa_xciv/ocr-3-device-uat-gate"]], [1.7, 1.5, 3.3])
    add_callout(doc, "Scope boundary", "This record documents enhancements to the running Senatla Ops application. It is not a plan for a replacement application. Harmony/Saaiplass material remains illustrative until contractually validated.", "FFF8E1")
    doc.add_page_break()

    doc.add_heading("Contents", level=1)
    for entry in (
        "1. Purpose and operating principles",
        "2. Existing product baseline",
        "3. Enhancement map",
        "4. Slice register",
        "5. Delivery sequence and first milestone",
        "6. Journey log",
        "7. Decision log",
        "8. Evidence register",
        "9. Risk register",
        "10. Definition of done",
        "11. Source material",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(entry)
    doc.add_paragraph("This static contents list remains readable in all renderers. Update it when a future revision adds or renames a major section.")
    doc.add_page_break()

    doc.add_heading("1. Purpose and operating principles", level=1)
    doc.add_paragraph("This living record connects each Senatla Ops implementation slice to the existing product, operational need, decisions, changes, verification evidence, risks, and observed outcomes. It must be updated after every meaningful implementation or verification pass.")
    add_callout(doc, "Source interpretation", "The supplied Operations Research document is a source of enhancement ideas only. Its scenario, project code, quantities, activities, and commercial assumptions are not production facts.")
    principles = [
        "Enhance current workflows before adding new modules.",
        "Reuse Supabase RLS, immutable audit events, private evidence storage, and the offline outbox.",
        "Treat safety, legal, technical, and contractual requirements as hard constraints.",
        "Keep recommendations explainable and subject to human approval.",
        "Derive dashboards from persisted records and preserve one source of truth.",
        "Require measured evidence before adding solver-backed optimisation.",
        "Separate verified completion from assumptions, partial proof, and external blockers.",
    ]
    for item in principles:
        add_bullet(doc, item)

    doc.add_heading("2. Existing product baseline", level=1)
    doc.add_paragraph("The document starts from the system that is already running. The following capabilities are foundations to reuse, not work to recreate.")
    baseline = [
        "Supabase authentication, application roles, route guards, and row-level security.",
        "Sites, teams, employees, attendance, safety talks, issues, and daily synchronisation.",
        "Employee onboarding, medical status, red-ticket fields, and PPE records.",
        "Asset identity, evidence, custody, compliance, maintenance, meters, and fuel.",
        "Vendor accounts, vendor invoices, Director review, and maker-checker approvals.",
        "Immutable administrative and attendance audit events.",
        "Durable outbox and queued offline submissions.",
        "Site Manager, Office Admin, Asset Register, and Director experiences.",
        "Director metrics based on recorded attendance, PPE, fuel, and vendor data.",
    ]
    for item in baseline:
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("3. Enhancement map", level=1)
    add_table(doc, ["Current feature", "Enhancement", "Boundary"], [
        ["Site Manager daily setup", "Crew, asset, and site readiness", "No replacement field app"],
        ["Employee onboarding", "Eligibility explanations", "No duplicate workforce register"],
        ["Asset control", "Availability and compliance readiness", "No duplicate asset register"],
        ["Office Admin", "Conflicts and eligible alternatives", "No automatic reassignment"],
        ["Trip logging", "Stop order and actual comparison", "No early optimality claim"],
        ["Cost records", "Site/job attribution", "No new accounting ledger"],
        ["Director dashboard", "Readiness, utilisation, blockers, cost", "No manual KPI totals"],
        ["Audit/outbox", "Recommendation and outcome evidence", "No parallel event store"],
    ], [1.55, 3.05, 1.9])

    doc.add_heading("4. Slice register", level=1)
    slices = [
        ("Slice 0 - Readiness contract", "Define deterministic ready, warning, blocked, and unknown outcomes with stable reason codes, corrective actions, fail-closed safety behavior, and a versioned policy.", "Rule catalogue; evaluator tests; privacy review; no duplicate tables.", "Implemented locally; frontend policy verified; database runtime proof pending local Docker/Supabase"),
        ("Slice 1 - Site Manager readiness", "Show current crew and asset readiness inside the existing daily setup; confirm current policy evidence atomically through the immutable audit contract.", "Component, privacy, real-data boundary, role, desktop, and 390 px evidence.", "Implemented and locally verified; remote database/deployment gate pending"),
        ("Slice 2 - Office Admin suggestions", "Detect conflicts, rank eligible alternatives transparently, and require human acceptance, rejection, or override reason.", "Ranking, conflict, blocker, audit, and role-negative tests.", "Implemented and locally verified; 30-assertion database runtime proof and remote deployment pending"),
        ("Slice 3 - Site/job cost attribution", "Attribute labour, PPE, fuel, work-order, and approved vendor costs to existing sites/job numbers; expose unattributed costs.", "Cost reconciliation, allocation, invoice-state, RLS, and role tests.", "Planned; after vendor baseline stabilises"),
        ("Slice 4 - Director control view", "Add drillable readiness, utilisation, blocker, cost, freshness, and completeness metrics to the existing Director page.", "Metric reconciliation, partial-data, access, desktop, and mobile evidence.", "Planned; depends on Slices 1 and 3"),
        ("Slice 5 - Trip recommendation", "Use current trip data to recommend a deterministic stop order and compare estimated with actual execution.", "Algorithm, coordinate, offline, audit, and mobile tests.", "Planned; depends on Slice 1"),
        ("Slice 6 - Optimisation evidence gate", "Measure real complexity and benefit, then decide whether deterministic rules remain sufficient or a solver is justified.", "Pilot dataset, measured baseline, decision record, and approved architecture if needed.", "Planned; after Slice 5 evidence"),
    ]
    for title, objective, evidence, status in slices:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        r = p.add_run("Objective. ")
        r.bold = True
        p.add_run(objective)
        p = doc.add_paragraph()
        r = p.add_run("Exit evidence. ")
        r.bold = True
        p.add_run(evidence)
        p = doc.add_paragraph()
        r = p.add_run("Status. ")
        r.bold = True
        p.add_run(status)

    doc.add_heading("5. Delivery sequence and first milestone", level=1)
    for item in [
        "Readiness rules.", "Site Manager readiness.", "Office Admin suggestions.",
        "Site/job cost attribution.", "Director operational control.",
        "Deterministic trip recommendation.", "Solver evidence decision.",
    ]:
        add_number(doc, item)
    add_callout(doc, "First milestone", "An existing Site Manager selects the current site, sees whether assigned people and assets are ready, understands every blocker, completes the existing safety and attendance workflow, and leaves an auditable readiness record.", "FFF8E1")

    doc.add_heading("6. Journey log", level=1)
    doc.add_paragraph("Add an entry after every meaningful implementation or verification pass. Copy the template below and preserve previous entries.")
    add_table(doc, ["Field", "Entry"], [
        ["Date", ""], ["Slice", ""], ["Branch / issue / PR", ""], ["Outcome sought", ""],
        ["Baseline before change", ""], ["Decision and rationale", ""], ["Alternative rejected and why", ""],
        ["Files / data contracts changed", ""], ["Security and privacy impact", ""], ["Tests executed", ""],
        ["Browser / mobile evidence", ""], ["Result", "Verified / Partial / Blocked"],
        ["Known gaps", ""], ["Next action", ""],
    ], [1.75, 4.75])
    doc.add_heading("23 August 2026 - Planning baseline", level=2)
    entries = [
        ("Slice", "Programme baseline"),
        ("Branch", "sankofa_xciv/admin-grouping-vendor-invoices"),
        ("Outcome", "Reframe the Operations Research source as enhancements to the running system"),
        ("Decision", "Preserve existing role experiences, stores, audit paths, and commercial workflows; begin with operational readiness"),
        ("Rejected alternative", "A separate project-management or optimisation platform before operational data proves the need"),
        ("Evidence", "Architecture, migrations, models, routes, role experiences, tests, and delivery backlog inspected"),
        ("Result", "Planning complete; implementation not started"),
        ("Known gap", "PDF text and page rendering completed; Codex image viewing was blocked by the managed Windows ACL helper"),
        ("Next action", "Implement and verify Slice 0"),
    ]
    add_table(doc, ["Field", "Recorded baseline"], entries, [1.55, 4.95])

    doc.add_heading("23 August 2026 - Slice 0 readiness foundation", level=2)
    slice_zero_entries = [
        ("Slice", "Slice 0 - Readiness contract and baseline"),
        ("Branch", "sankofa_xciv/admin-grouping-vendor-invoices"),
        ("Outcome", "Versioned readiness policy, sanitized site-readiness RPC, and five deterministic audit users with cross-site fixtures"),
        ("Decision", "Keep confidential evidence behind existing RLS and expose only outcomes, reason codes, corrective actions, policy version, and evaluation time"),
        ("Rejected alternative", "Direct Site Manager access to protected medical, screening, ticket, or onboarding records"),
        ("Changed contracts", "Readiness TypeScript module; one migration; local seed; two pgTAP contracts"),
        ("Security", "Public RPC is security-invoker; private evaluator pins an empty search path, validates session, organisation, role, and site access, and returns no raw onboarding fields"),
        ("Three-pass challenge", "Pass 1 found client/server rule drift; Pass 2 found an exposed privileged function and unsupported screening decisions; Pass 3 challenged determinism, missing evidence, work orders, deployability, and artifacts"),
        ("Improvements", "Aligned policy versions; removed unsupported screening decisions; added fail-closed evidence and invalid-date rules, renewal warnings, work-order blockers, private privileged evaluation, and an invoker API wrapper"),
        ("Tests", "TypeScript compilation passed; 12/12 focused policy tests passed; 36 RPC and 22 fixture database assertions authored"),
        ("Visual evidence", "Not applicable to this non-UI slice; full visual regression remains scheduled after Slice 3"),
        ("Result", "Partial verification: implementation and frontend policy verified; database runtime blocked pending Docker Linux engine"),
        ("Known gap", "Unavailable schema evidence remains unknown; no induction, skills, or medical-expiry rules were invented"),
        ("Next action", "Implement Slice 1, execute database contracts when Docker is available, run its hardened gate, then deploy Slices 0 and 1 together to remote dev; keep shared-password seed local only"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], slice_zero_entries, [1.55, 4.95])
    doc.add_heading("23 August 2026 - Slice 1 Site Manager live readiness", level=2)
    slice_one_entries = [
        ("Slice", "Slice 1 - Site Manager readiness panel"),
        ("Branch", "sankofa_xciv/admin-grouping-vendor-invoices"),
        ("Outcome", "Existing start-of-shift modal extended with authenticated site selection, sanitized readiness, corrective actions, fail-closed progression, and atomic server confirmation"),
        ("Decision", "Use the existing Supabase client through a narrow readiness gateway; only authenticated and RLS-filtered site UUIDs can drive readiness, sync, and audit"),
        ("Rejected alternative", "Broad Office Admin snapshot loading or client-side readiness heuristics that can fall back to demo records or expose excess evidence"),
        ("Changed contracts", "Readiness gateway/service; Site Manager component/template/tests; remote fallback guard; runtime loader; readiness migration/pgTAP; Cypress responsive check"),
        ("Security", "No demo fallback or remote autosave after missing snapshots; schema-validated safe labels/outcomes/actions; atomic re-evaluation; one sanitized immutable event per actor/site/day"),
        ("Three-pass challenge", "Pass 1 challenged identity, malformed responses, privacy, and fallback; Pass 2 found the refresh/confirm race; Pass 3 challenged startup, mobile reachability, duplication, and deployment"),
        ("Improvements", "Removed hard-coded Site Manager demo IDs; fail closed for loading, unavailable, blocked, unknown, malformed, and version mismatch; preserve GPS, targets, trip, safety, attendance, and sync; remove silent config fallback"),
        ("Tests", "Typecheck and build pass; targeted lint pass; 22/22 focused Chrome Headless tests; 3/3 runtime-config tests; 48 pgTAP assertions authored pending database execution"),
        ("Visual evidence", "Focused Cypress passed at desktop and 390 px with test-only intercepted responses; no mock data persisted; full regression remains scheduled after Slice 3"),
        ("Credentials", "Five-user register will be created only after real remote-dev provisioning in ignored output/credentials; no passwords or service keys in this DOCX or Git"),
        ("Result", "Local implementation and clean integration onto current origin/dev verified at release-candidate commit 41c8e62; remote deployment and provisioning blocked by target configuration"),
        ("Known gap", "Generated Supabase URL/key references differ; Vercel Preview Senatla values resolve empty; authenticated Supabase CLI lists no Senatla project"),
        ("Next action", "Link authoritative remote dev; populate matching Preview runtime values; execute contracts/migration; deploy; invite five real users; create controlled credential handoff; verify cross-role and cross-site audit flow"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], slice_one_entries, [1.55, 4.95])
    doc.add_heading("23 August 2026 - Slice 2 Office Admin reviewed assignments", level=2)
    slice_two_entries = [
        ("Slice", "Slice 2 - Office Admin conflicts and suggestions"),
        ("Branch", "sankofa_xciv/slice-01-readiness-release"),
        ("Outcome", "Existing employee bulk-site and asset custody workflows now include fail-closed reviews, ranked alternatives, controlled decisions, and atomic Supabase RPCs"),
        ("Decision", "Keep browser recommendations advisory and re-evaluate role, organisation, target, resource state, readiness, compliance, and work orders inside the mutation transaction"),
        ("Rejected alternative", "Client-only validation, local mutation, a new optimisation service, or an invented utilisation score without authoritative evidence"),
        ("Changed contracts", "Office Admin and Asset Operations UI; OfficeAdminService; assignment planner; one migration; focused TypeScript/component tests; 30-assertion pgTAP contract"),
        ("Security", "Office-only invoker RPCs call private definer implementations; hard blockers cannot be overridden; direct writes are trigger-blocked; protected details stay out of generic audit; trigger markers are cleared immediately"),
        ("Three-pass challenge", "Pass 1 found false no-change warnings/custody; Pass 2 aligned browser and database decisions; Pass 3 found a reusable transaction-local trigger marker in an eight-file security diff scan"),
        ("Improvements", "No-change is informational; unchanged rows are not rewritten; unchanged assets create no custody event; missing licence fails closed; marker lifetime is limited to the guarded update"),
        ("Tests", "Typecheck passed; 14/14 focused Chrome Headless tests passed; 30 authorization, role, blocker, override, no-op, custody, direct-write, and leak assertions authored"),
        ("Visual evidence", "Existing responsive interfaces retain 44 px decision controls; full visual regression remains scheduled after Slice 3"),
        ("Result", "Partial: implementation, browser tests, and hardened security remediation pass; database execution blocked because Docker Desktop Linux engine is not running"),
        ("Known gap", "Recent utilisation is excluded until an authoritative metric exists; remote deployment and five-user provisioning await a verified Senatla remote-dev target"),
        ("Next action", "Commit the clean candidate; execute pgTAP and deploy only after remote target verification; then provision five real users and create the ignored credential handoff"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], slice_two_entries, [1.55, 4.95])
    doc.add_heading("24 August 2026 - OCR-0 native platform foundation", level=2)
    ocr_zero_entries = [
        ("Slice", "OCR-0 - production native backbone before OCR logic"),
        ("Branch / PR", "sankofa_xciv/authoritative-slices-0-2; PR #53 into dev"),
        ("Outcome", "Committed Android and iOS Capacitor projects under za.co.senatlatrading.ops, exact install-script approvals, unsigned native gates, and native privacy controls; no OCR SDK or replacement app"),
        ("Decision", "Prove the native security and build boundary before scanning contracts or Android ML Kit and Apple Vision/VisionKit adapters"),
        ("Rejected alternative", "Browser-only capture or an immediate cross-platform OCR dependency that would weaken native quality, offline behavior, privacy control, and platform ownership"),
        ("Changed contracts", "Capacitor identity/dependencies; android and ios projects; native CI; Swift path normalizer; browser baseline; signing exclusions; native guide"),
        ("Security", "Android backup disabled; FileProvider restricted to a dedicated evidence cache; iOS purposes explicit; signing/API keys ignored; no secret stored in Git or this document"),
        ("Three-pass challenge", "Pass 1 aligned IDs, versions, browsers, permissions, and sync; Pass 2 found backup/file-sharing leaks, stale test identities, secret exclusions, trigger gaps, and silent normalization; Pass 3 ran executable and remote gates"),
        ("Improvements", "Disabled backups; narrowed sharing; corrected native test packages; compiled Android tests in CI; expanded triggers; made iOS normalization fail closed; documented install-script trust and release gates"),
        ("Tests", "Runtime-config 3/3, lint, typecheck, production build, and 80/80 browser tests passed; repeated Android/iOS sync passed; production dependency audit found zero vulnerabilities"),
        ("Native evidence", "GitHub Actions run 32676423136 passed: Android app lint, unit tests, instrumentation-test APK compilation, and unsigned debug APK built on Linux in 3m54s; the unsigned iOS simulator app built on macOS in 2m14s. https://github.com/Sankofa-Digital-Studio/senatla-ops/actions/runs/32676423136"),
        ("Visual evidence", "No rendered app UI changed; full regression remains after Slice 3 and scanner visual evidence begins with the first rendered OCR slice"),
        ("Result", "Verified native foundation, subject to recorded release-stage gates"),
        ("Known gaps", "Local Android compile unavailable without JDK/SDK; iOS requires macOS; release signing, shrinking/version allocation, store identity, and Swift dependency lock remain release-stage work"),
        ("Next action", "OCR-1 adds the stable TypeScript scanning contract and privacy lifecycle before native OCR adapters"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], ocr_zero_entries, [1.55, 4.95])
    doc.add_heading("24 August 2026 - OCR-1 and OCR-2 production scanning bundle", level=2)
    ocr_one_two_entries = [
        ("Slice", "OCR-1 stable scan contract and OCR-2 Android/iOS native adapters"),
        ("Branch / PR", "sankofa_xciv/ocr-1-2-native-scanning; PR #54 merged into dev as 05df90a"),
        ("Outcome", "Existing asset registration enhanced with native scanning, on-device OCR, evidence provenance, byte-integrity verification, and explicit human review"),
        ("Decision", "Use a strict TypeScript/native contract; private JPEG artifacts only; validate dimensions, size, UUID and SHA-256; delete native artifacts after verified JavaScript ownership"),
        ("Rejected alternative", "Base64 bridge payloads, public native files, default cloud OCR, automatic field application, and silent fallback after denial/cancellation"),
        ("Security", "Bounded evidence and OCR text; no native paths or raw OCR in logs; private Storage/RLS authoritative; human apply-and-verify required"),
        ("Three-pass challenge", "Challenged contract/fallback/integrity, native lifecycle and cleanup, then all changed source and RLS/Storage controls; persistence findings were remediated before closure"),
        ("Improvements", "Held Android guard through finalisation; safe startup sweeps; symlink rejection; main-thread bridge resolution; whole-session cleanup; checked rollback; immutable ready evidence"),
        ("Tests", "Local release gate and 93/93 browser tests passed; PR #54 passed app, 208 pgTAP tests, Android, iOS and Vercel checks"),
        ("Result", "Implementation, migrations, native CI and deployments passed; signed physical-device proof remains the store-release gate"),
        ("Known gaps", "OCR accuracy and permission/lifecycle behavior on representative physical devices require the five-role signed-device UAT matrix"),
        ("Next action", "Execute OCR-3 private signed-device UAT without placing credentials or direct identifiers in Git or this document"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], ocr_one_two_entries, [1.55, 4.95])
    doc.add_heading("24 August 2026 - OCR-3 signed-device UAT gate", level=2)
    ocr_three_entries = [
        ("Slice", "OCR-3 - five-role signed-device release gate"),
        ("Branch", "sankofa_xciv/ocr-3-device-uat-gate"),
        ("Outcome", "Executable contract for five real access variants, signed Android/iOS devices, and 25 required results; private ignored evidence boundary and responsive guide"),
        ("Decision", "Treat physical UAT as release evidence, not seed data; retain aliases and safe evidence references only; bind proof to exact remote target and deployed commit"),
        ("Rejected alternative", "Committed credentials/results, document-only attestation, synthetic users presented as UAT, or a single happy-path account"),
        ("Security", "Reject credentials, direct identifiers, raw OCR, native/private paths and device identifiers recursively; output summary/issues only, never the submitted payload or caller path"),
        ("Three-pass challenge", "Pass 1 reconciled 16 scenarios to 25 results; Pass 2 challenged leakage, evidence traversal and unknown additions; Pass 3 challenged signed-build, commit and release semantics"),
        ("Improvements", "Corrected result count; reject unknown platforms/cases; isolate private JSON; incomplete mode can validate structure but cannot mark release ready"),
        ("Tests", "6/6 contract tests pass: complete matrix, unknown additions, leakage, missing/failed evidence, and in-progress behavior"),
        ("Visual evidence", "Responsive guide supports Android/iOS viewing; application UI is unchanged; physical scanner proof remains pending"),
        ("Result", "Gate verified locally; no UAT outcome fabricated and OCR is not yet device-certified"),
        ("Known gaps", "Signed builds, five real assignments, controlled credential handoff, and physical evidence require authorised operators"),
        ("Next action", "Deploy after full CI, execute controlled physical UAT, and validate the ignored private record"),
    ]
    add_table(doc, ["Field", "Recorded implementation evidence"], ocr_three_entries, [1.55, 4.95])
    doc.add_heading("7. Decision log", level=1)
    add_table(doc, ["ID", "Date", "Decision", "Rationale", "Status"], [
        ["D-001", "23 Aug 2026", "Enhance existing app", "Foundation already exists", "Accepted"],
        ["D-002", "23 Aug 2026", "Begin with readiness", "Immediate safety and operational value", "Accepted"],
        ["D-003", "23 Aug 2026", "Human-reviewed recommendations", "Preserve accountability", "Accepted"],
        ["D-004", "23 Aug 2026", "Defer solver", "Measure complexity and benefit first", "Accepted"],
        ["D-005", "23 Aug 2026", "Sanitize readiness at database boundary", "Site users need outcomes, not confidential evidence", "Accepted"],
        ["D-006", "23 Aug 2026", "Bundle Slices 0 and 1 for remote dev", "Deploy the contract with its verified consumer", "Accepted"],
        ["D-007", "23 Aug 2026", "No remote demo fallback", "Missing reads cannot create operational records", "Accepted"],
        ["D-008", "23 Aug 2026", "Atomic server confirmation", "Prevent stale UI self-certification", "Accepted"],
        ["D-009", "23 Aug 2026", "Credentials outside Git and DOCX", "Controlled UAT handoff", "Accepted"],
        ["D-010", "23 Aug 2026", "Atomic assignment re-evaluation and short-lived trigger marker", "Prevent stale decisions and direct-write bypass", "Accepted"],
        ["D-011", "24 Aug 2026", "Native foundation before OCR logic", "Control identity, privacy, and build evidence first", "Accepted"],
        ["D-012", "24 Aug 2026", "Native on-device OCR behind strict contract", "Private bounded evidence and offline capability", "Accepted"],
        ["D-013", "24 Aug 2026", "OCR remains human-reviewed evidence", "Prevent recognition errors becoming authoritative data", "Accepted"],
        ["D-014", "24 Aug 2026", "Private alias-only executable physical UAT", "Prevent false certification and sensitive-data leakage", "Accepted"],
    ], [0.6, 0.9, 1.6, 2.45, 0.95])

    doc.add_heading("8. Evidence register", level=1)
    rows = []
    for i, evidence in enumerate([
        "Readiness rule catalogue and evaluator tests",
        "Site Manager privacy, offline, component, and browser proof",
        "Conflict, ranking, blocker, and audit tests",
        "Cost reconciliation and RLS tests",
        "KPI reconciliation and responsive browser proof",
        "Trip heuristic, offline, audit, and field proof",
        "Pilot measures and solver decision",
    ], start=1):
        if i == 1:
            rows.append([f"E-{i:03}", str(i - 1), evidence, "Partial", "12/12 policy tests pass; 48 RPC assertions await database execution"])
        elif i == 2:
            rows.append([f"E-{i:03}", str(i - 1), evidence, "Partial", "22/22 focused tests, build, lint, and desktop/390 px pass; remote proof pending"])
        elif i == 3:
            rows.append([f"E-{i:03}", str(i - 1), evidence, "Partial", "14/14 focused tests and security remediation pass; 30 pgTAP assertions await execution"])
        else:
            rows.append([f"E-{i:03}", str(i - 1), evidence, "Pending", "To be recorded"])
    add_table(doc, ["ID", "Slice", "Evidence", "Status", "Reference"], rows, [0.7, 0.55, 3.0, 0.85, 1.4])

    add_table(doc, ["ID", "Slice", "Evidence", "Status", "Reference"], [
        ["E-008", "OCR-0", "Android/iOS identity, sync, security, and unsigned build proof", "Verified", "PR #53 native CI recorded in journey log"],
        ["E-009", "OCR-1", "Contract, integrity, cleanup, fallback, and review proof", "Verified", "Release gate, security scan, and browser tests passed"],
        ["E-010", "OCR-2", "Android ML Kit and iOS VisionKit/Vision compile and device proof", "Partial", "Remote native CI passed; signed physical-device UAT pending"],
        ["E-011", "OCR-3", "Five-role UAT contract, privacy gate, and execution guide", "Verified locally", "6/6 contract tests pass; physical results pending"],
    ], [0.7, 0.75, 2.75, 0.85, 1.45])

    doc.add_heading("9. Risk register", level=1)
    add_table(doc, ["Risk", "Consequence", "Control", "State"], [
        ["Illustrative contract assumptions treated as fact", "Contractual or reputational harm", "Validate customer-specific requirements", "Open"],
        ["Sensitive screening detail exposed", "Privacy breach", "Least-detail, role-filtered reasons", "Open"],
        ["UI-only safety checks", "Unsafe allocation", "Authoritative server/database rules", "Open"],
        ["Duplicate operational structures", "Divergent truth", "Extend existing entities and gateways", "Controlled"],
        ["Dashboard drift", "Incorrect decisions", "Drill-down and reconciliation tests", "Open"],
        ["Recommendation mistaken for instruction", "Safety/accountability risk", "Human decision and override audit", "Controlled"],
        ["Premature solver", "Cost without value", "Production evidence gate", "Controlled"],
        ["Interrupted scan leaves cache residue", "Private evidence persists", "Private cache, cleanup and startup sweep", "Code controlled; device test pending"],
        ["OCR suggestion treated as truth", "Incorrect asset data", "Human apply/verify and immutable provenance", "Code controlled; UAT pending"],
    ], [1.65, 1.45, 2.45, 0.95])

    doc.add_heading("10. Definition of done", level=1)
    for item in [
        "The existing workflow is enhanced rather than duplicated.",
        "Acceptance criteria are demonstrated, not only described.",
        "New mutations leave immutable audit evidence.",
        "Authorization changes include role-negative tests.",
        "Schema changes include migration/reset and rollback evidence.",
        "Rendered changes have desktop and 390 px proof.",
        "Metrics reconcile to persisted records.",
        "The journey, decision, evidence, and risk registers are updated.",
        "Verified completion is separated from partial proof and external blockers.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("11. Source material", level=1)
    doc.add_paragraph("Senatla Trading - Harmony Mine Operations: Saaiplass #3 Rehabilitation & Demolition - Discrete Programming + Operations Research Business Ecosystem. Thirteen-page illustrative strategy document supplied August 2026.")
    doc.add_paragraph("Current Senatla Ops architecture, delivery backlog, migrations, models, routes, role experiences, and tests inspected during the planning pass.")

    props = doc.core_properties
    props.title = "Senatla Ops Implementation Journey"
    props.subject = "Living record for evidence-led enhancement of the existing Senatla Ops application"
    props.author = "Senatla Trading / Sankofa Digital"
    props.keywords = "Senatla Ops, implementation journey, readiness, operational control"
    props.comments = "Living working document - update after every implementation slice."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
