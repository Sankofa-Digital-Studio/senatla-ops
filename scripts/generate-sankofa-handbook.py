from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "sankofa-senatla-ops-engineering-handbook.md"
HTML_OUT = ROOT / "docs" / "sankofa-senatla-ops-engineering-handbook.html"
PDF_OUT = ROOT / "output" / "pdf" / "Sankofa-Senatla-Ops-Engineering-Handbook.pdf"
DOCX_OUT = ROOT / "output" / "doc" / "Sankofa-Senatla-Ops-Engineering-Handbook.docx"

NAVY = colors.HexColor("#0B172A")
BLUE = colors.HexColor("#1E5AA8")
GOLD = colors.HexColor("#E4B341")
LIGHT = colors.HexColor("#F3F6FA")
MID = colors.HexColor("#536172")


def clean(value: str) -> str:
    return value.replace("—", "-").replace("–", "-").replace("‑", "-").strip()


def blocks(text: str):
    lines = [clean(line.rstrip()) for line in text.splitlines()]
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line:
            index += 1
            continue
        if line.startswith("```"):
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                index += 1
            index += 1
            continue
        image_match = re.match(r"^!\[(.+?)\]\((.+?)\)$", line)
        if image_match:
            yield "image", (image_match.group(1), image_match.group(2))
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            table_lines = [line]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [[clean(cell) for cell in row.strip("|").split("|")] for row in table_lines]
            yield "table", rows
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            yield "heading", (len(heading.group(1)), heading.group(2))
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            items = []
            while index < len(lines) and re.match(r"^[-*]\s+", lines[index]):
                items.append(re.sub(r"^[-*]\s+", "", lines[index]))
                index += 1
            yield "list", items
            continue
        paragraph = [line]
        index += 1
        while index < len(lines) and lines[index] and not lines[index].startswith(("#", "|", "- ", "* ")):
            paragraph.append(lines[index])
            index += 1
        yield "paragraph", " ".join(paragraph)


def inline_markup(value: str) -> str:
    escaped = html.escape(clean(value))
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`(.+?)`", r"<font name='Courier'>\1</font>", escaped)
    escaped = re.sub(r"(https?://[^\s<]+)", r"<link href='\1' color='#1E5AA8'>\1</link>", escaped)
    return escaped


class HandbookDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, footer_label: str):
        self.footer_label = footer_label
        super().__init__(filename, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=22 * mm, bottomMargin=20 * mm)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body")
        self.addPageTemplates(PageTemplate(id="content", frames=[frame], onPage=self.draw_page))

    def draw_page(self, canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D6DEE8"))
        canvas.line(18 * mm, A4[1] - 15 * mm, A4[0] - 18 * mm, A4[1] - 15 * mm)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(NAVY)
        canvas.drawString(18 * mm, A4[1] - 11 * mm, "SANKOFA | SENATLA OPS")
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(MID)
        canvas.drawRightString(A4[0] - 18 * mm, A4[1] - 11 * mm, "CONFIDENTIAL - INTERNAL USE")
        canvas.drawString(18 * mm, 10 * mm, self.footer_label)
        canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, f"Page {doc.page}")
        canvas.restoreState()

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            level = getattr(flowable, "toc_level", None)
            if level is not None:
                key = f"section-{self.seq.nextf('heading')}"
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(flowable.getPlainText(), key, level=level, closed=False)
                self.notify("TOCEntry", (level, flowable.getPlainText(), self.page, key))


def build_pdf(parsed, output=PDF_OUT, document_title="Sankofa Senatla Ops", document_subtitle="Engineering Handbook", description="Controlled release baseline for Senatla Trading's internal operational products", footer_label="Engineering Handbook v1.0", cover_image=None):
    output.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=30, textColor=NAVY, alignment=TA_LEFT, spaceAfter=12)
    subtitle = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=11, leading=16, textColor=MID)
    heading1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, leading=20, textColor=NAVY, spaceBefore=14, spaceAfter=7)
    heading2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, leading=15, textColor=BLUE, spaceBefore=10, spaceAfter=5)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=8.8, leading=12.2, textColor=colors.HexColor("#253247"), spaceAfter=5)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=12, firstLineIndent=-7, bulletIndent=3)

    story = [Spacer(1, 20 * mm), Paragraph(document_title, title), Paragraph(document_subtitle, ParagraphStyle("CoverSub", parent=title, fontSize=18, textColor=BLUE)), Spacer(1, 6 * mm), Paragraph(description, subtitle)]
    if cover_image and Path(cover_image).exists():
        cover_visual = RLImage(str(cover_image))
        scale = min((165 * mm) / cover_visual.imageWidth, (62 * mm) / cover_visual.imageHeight)
        cover_visual.drawWidth = cover_visual.imageWidth * scale
        cover_visual.drawHeight = cover_visual.imageHeight * scale
        story.extend([Spacer(1, 7 * mm), cover_visual, Spacer(1, 7 * mm)])
    else:
        story.append(Spacer(1, 18 * mm))
    cover = Table([["VERSION", "1.0"], ["DATE", "23 June 2026"], ["STATUS", "Controlled release baseline"], ["CLASSIFICATION", "Confidential - internal use"]], colWidths=[42 * mm, 105 * mm])
    cover.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, -1), NAVY), ("TEXTCOLOR", (0, 0), (0, -1), colors.white), ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTNAME", (1, 0), (1, -1), "Helvetica"), ("FONTSIZE", (0, 0), (-1, -1), 9), ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CAD4E0")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([cover, Spacer(1, 12 * mm), Paragraph("Prepared by Sankofa Digital Studio for Senatla Trading", subtitle), PageBreak()])
    story.append(Paragraph("Table of contents", title))
    toc = TableOfContents()
    toc.levelStyles = [ParagraphStyle("TOC1", fontName="Helvetica", fontSize=10, leading=15, leftIndent=0, firstLineIndent=0, textColor=NAVY), ParagraphStyle("TOC2", fontName="Helvetica", fontSize=9, leading=13, leftIndent=12, firstLineIndent=0, textColor=MID)]
    story.extend([toc, PageBreak()])

    first_heading = True
    for kind, value in parsed:
        if kind == "heading":
            level, text = value
            if first_heading:
                first_heading = False
                continue
            style = heading1 if level <= 2 else heading2
            item = Paragraph(inline_markup(text), style)
            item.toc_level = 0 if level <= 2 else 1
            story.append(item)
        elif kind == "paragraph":
            story.append(Paragraph(inline_markup(value), body))
        elif kind == "list":
            for item in value:
                story.append(Paragraph(inline_markup(item), bullet, bulletText="-"))
        elif kind == "image":
            alt, relative_path = value
            image_path = SOURCE.parent / relative_path
            if image_path.exists():
                image = RLImage(str(image_path))
                scale = min((174 * mm) / image.imageWidth, (205 * mm) / image.imageHeight)
                image.drawWidth = image.imageWidth * scale
                image.drawHeight = image.imageHeight * scale
                story.extend([image, Paragraph(inline_markup(alt), ParagraphStyle("Caption", parent=body, fontSize=8, textColor=MID, alignment=TA_CENTER)), Spacer(1, 8)])
        elif kind == "table":
            data = [[Paragraph(inline_markup(cell), body) for cell in row] for row in value]
            col_width = 174 * mm / max(1, len(data[0]))
            table = Table(data, colWidths=[col_width] * len(data[0]), repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7D1DD")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
            story.extend([table, Spacer(1, 6)])

    doc = HandbookDocTemplate(str(output), footer_label)
    doc.multiBuild(story)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, code, separate, end])


def build_docx(parsed, output=DOCX_OUT, document_title="Sankofa Senatla Ops", document_subtitle="Engineering Handbook", description="Controlled release baseline for Senatla Trading's internal operational products", footer_label="Engineering Handbook v1.0"):
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    for name, size, color in [("Title", 28, "0B172A"), ("Heading 1", 17, "0B172A"), ("Heading 2", 13, "1E5AA8")]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "SANKOFA | SENATLA OPS                                      CONFIDENTIAL - INTERNAL USE"
    header.style = styles["Caption"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"{footer_label} | ")
    add_field(footer, "PAGE")

    p = doc.add_paragraph()
    p.add_run(document_title).bold = True
    p.style = styles["Title"]
    sub = doc.add_paragraph(document_subtitle)
    sub.style = styles["Subtitle"]
    doc.add_paragraph(description)
    doc.add_paragraph("Version 1.0 | 23 June 2026 | Confidential - internal use")
    doc.add_page_break()
    doc.add_heading("Table of contents", level=1)
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    first_heading = True
    for kind, value in parsed:
        if kind == "heading":
            level, text = value
            if first_heading:
                first_heading = False
                continue
            doc.add_heading(clean(text), level=min(3, max(1, level - 1)))
        elif kind == "paragraph":
            doc.add_paragraph(re.sub(r"\*\*|`", "", clean(value)))
        elif kind == "list":
            for item in value:
                doc.add_paragraph(re.sub(r"\*\*|`", "", clean(item)), style="List Bullet")
        elif kind == "image":
            alt, relative_path = value
            image_path = SOURCE.parent / relative_path
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.7))
                caption = doc.add_paragraph(alt)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.style = styles["Caption"]
        elif kind == "table":
            table = doc.add_table(rows=len(value), cols=len(value[0]))
            table.style = "Light Shading Accent 1"
            for row_index, row in enumerate(value):
                for column_index, cell in enumerate(row):
                    table.cell(row_index, column_index).text = re.sub(r"\*\*|`", "", clean(cell))
                    if row_index == 0:
                        for run in table.cell(row_index, column_index).paragraphs[0].runs:
                            run.bold = True
            doc.add_paragraph()
    doc.save(output)


def build_html(parsed, output=HTML_OUT, page_title="Sankofa Senatla Ops Engineering Handbook", hero_title="Senatla Ops<br>Engineering Handbook", description="Controlled release baseline for Senatla Trading's internal operational products."):
    nav = []
    body_parts = []
    first_heading = True
    for kind, value in parsed:
        if kind == "heading":
            level, text = value
            if first_heading:
                first_heading = False
                continue
            anchor = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
            if level <= 2:
                nav.append(f'<a href="#{anchor}">{html.escape(text)}</a>')
            body_parts.append(f'<h{min(level, 4)} id="{anchor}">{html.escape(text)}</h{min(level, 4)}>')
        elif kind == "paragraph":
            body_parts.append(f"<p>{inline_markup(value)}</p>")
        elif kind == "list":
            body_parts.append("<ul>" + "".join(f"<li>{inline_markup(item)}</li>" for item in value) + "</ul>")
        elif kind == "image":
            alt, relative_path = value
            body_parts.append(f'<figure><img src="{html.escape(relative_path)}" alt="{html.escape(alt)}"><figcaption>{html.escape(alt)}</figcaption></figure>')
        elif kind == "table":
            head = "".join(f"<th>{inline_markup(cell)}</th>" for cell in value[0])
            rows = "".join("<tr>" + "".join(f"<td>{inline_markup(cell)}</td>" for cell in row) + "</tr>" for row in value[1:])
            body_parts.append(f'<div class="table-wrap"><table><thead><tr>{head}</tr></thead><tbody>{rows}</tbody></table></div>')
    document = f'''<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1,viewport-fit=cover"><title>{html.escape(page_title)}</title>
<style>:root{{--navy:#0b172a;--blue:#1e5aa8;--gold:#e4b341;--paper:#f4f7fb;--ink:#253247;--line:#d4deea}}*{{box-sizing:border-box}}html{{scroll-behavior:smooth}}body{{margin:0;background:var(--paper);color:var(--ink);font:16px/1.65 system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}}header{{padding:max(2rem,env(safe-area-inset-top)) max(1rem,env(safe-area-inset-right)) 2.5rem max(1rem,env(safe-area-inset-left));background:linear-gradient(135deg,var(--navy),#173b6c);color:white}}header .inner,main{{max-width:1120px;margin:auto}}.eyebrow{{color:var(--gold);font-weight:800;letter-spacing:.14em;text-transform:uppercase;font-size:.75rem}}h1{{font-size:clamp(2rem,7vw,4.5rem);line-height:1.05;margin:.6rem 0}}header p{{max-width:720px;color:#dce8f7}}.meta{{display:flex;flex-wrap:wrap;gap:.65rem;margin-top:1.5rem}}.meta span{{border:1px solid #ffffff40;border-radius:999px;padding:.35rem .75rem;font-size:.8rem}}main{{display:grid;grid-template-columns:240px minmax(0,1fr);gap:2rem;padding:2rem 1rem 5rem}}nav{{position:sticky;top:1rem;align-self:start;background:white;border:1px solid var(--line);border-radius:16px;padding:1rem;max-height:calc(100vh - 2rem);overflow:auto}}nav a{{display:block;color:#40546f;text-decoration:none;padding:.35rem .5rem;border-radius:8px;font-size:.84rem}}nav a:hover,nav a:focus{{background:#eaf1fb;color:var(--blue)}}article{{min-width:0;background:white;border:1px solid var(--line);border-radius:20px;padding:clamp(1rem,4vw,3rem);box-shadow:0 18px 55px #18345412}}h2{{color:var(--navy);font-size:1.6rem;margin:2.4rem 0 .8rem;border-top:1px solid var(--line);padding-top:1.3rem}}h3{{color:var(--blue)}}p,li{{max-width:78ch}}figure{{margin:1.5rem 0}}figure img{{width:100%;height:auto;border:1px solid var(--line);border-radius:14px;display:block}}figcaption{{font-size:.82rem;color:#617086;margin-top:.45rem;text-align:center}}.table-wrap{{overflow-x:auto;margin:1rem 0;border:1px solid var(--line);border-radius:12px}}table{{border-collapse:collapse;width:100%;min-width:620px;font-size:.9rem}}th{{background:var(--navy);color:white;text-align:left}}th,td{{padding:.7rem;border-bottom:1px solid var(--line);vertical-align:top}}tbody tr:nth-child(even){{background:#f7f9fc}}a{{color:var(--blue);overflow-wrap:anywhere}}@media(max-width:800px){{main{{display:block;padding:.75rem .75rem 4rem}}nav{{position:relative;top:0;max-height:none;margin-bottom:.75rem}}article{{border-radius:14px;padding:1rem}}h2{{font-size:1.35rem}}header{{padding-bottom:1.8rem}}}}@media print{{nav{{display:none}}main{{display:block;padding:0}}article{{border:0;box-shadow:none}}header{{background:white;color:var(--navy);padding:1rem 0}}header p{{color:var(--ink)}}}}</style></head>
<body><header><div class="inner"><div class="eyebrow">Sankofa Digital Studio | Confidential</div><h1>{hero_title}</h1><p>{html.escape(description)}</p><div class="meta"><span>Version 1.0</span><span>23 June 2026</span><span>Internal use</span></div></div></header><main><nav aria-label="Table of contents">{''.join(nav)}</nav><article>{''.join(body_parts)}</article></main></body></html>'''
    output.write_text(document, encoding="utf-8")


if __name__ == "__main__":
    parsed = list(blocks(SOURCE.read_text(encoding="utf-8")))
    build_html(parsed)
    build_pdf(parsed)
    build_docx(parsed)
    print(HTML_OUT)
    print(PDF_OUT)
    print(DOCX_OUT)
